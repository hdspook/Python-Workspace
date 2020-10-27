import os 

PROJECT_DIR=os.path.abspath('.')
OUTPUT_FILE_URL = os.path.join(PROJECT_DIR,'static\splitter\output')


def create(data):
    import docx 

    # create an instance of a word document 
    doc = docx.Document()  
    doc.add_heading('Patient Details', 0) 

    # add a paragraph and store 
    # the object in a variable
    doc_para = doc.add_paragraph('XYZ Hospital \n\n ') 
    for details in data:
        p = doc_para.add_run('Details of patient with id ' + str(details['id']) + '\n')
        p.italic = True
        p.bold = True
        p.underline = True
        doc_para.add_run('Name  : ' + str(details['name']) + '\n').italic = True
        doc_para.add_run('Email : ' + str(details['email']) + '\n').italic = True
        doc_para.add_run('Age   : ' + str(details['age']) + '\n').italic = True
        doc_para.add_run('\n\n\n')

    # add a page break to start a new page 
    doc.add_page_break() 

    # now save the document to a location 
    doc.save(OUTPUT_FILE_URL + "\\test.docx") 
